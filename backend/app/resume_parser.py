import os
import re

from docx import Document
import PyPDF2


def clean_text(text: str) -> str:
    """
    Clean extracted resume text while preserving useful
    technology names such as C#, C++, .NET and ASP.NET.
    """
    if not text:
        return ""

    # Normalize common whitespace characters
    text = text.replace("\x00", " ")
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)

    # Remove excessive blank lines
    text = re.sub(r"\n\s*\n+", "\n", text)

    return text.strip()


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a normal text-based PDF.
    """
    text_parts = []

    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError(
                    "The uploaded PDF is password protected. "
                    "Please upload an unlocked PDF."
                )

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text()

                if page_text:
                    text_parts.append(page_text)

            except Exception as exc:
                # Continue processing other pages instead of
                # completely failing the resume.
                print(
                    f"Warning: Could not extract page "
                    f"{page_number}: {exc}"
                )

    return clean_text("\n".join(text_parts))


def extract_docx_text(file_path: str) -> str:
    """
    Extract text from DOCX paragraphs and tables.
    """
    doc = Document(file_path)

    text_parts = []

    # Normal paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Tables are very common in resumes.
    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                text_parts.append(" | ".join(row_text))

    return clean_text("\n".join(text_parts))


def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from PDF or DOCX resume.

    Raises ValueError when:
    - the file format is unsupported
    - the file contains no extractable text
    """

    if not file_path:
        raise ValueError("Resume file path is empty.")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        text = extract_pdf_text(file_path)

    elif extension == ".docx":
        text = extract_docx_text(file_path)

    else:
        raise ValueError(
            "Unsupported file format. "
            "Only PDF and DOCX files are supported."
        )

    # Important:
    # Do NOT allow an empty extracted resume to continue
    # to the scoring function and silently become 0.
    if not text or len(text.strip()) < 20:
        raise ValueError(
            "Could not extract readable text from this resume. "
            "If this is a scanned/image-based PDF, please upload "
            "a text-based PDF or DOCX file."
        )

    return text

